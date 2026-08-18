import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_icact_docx(md_path, docx_path):
    doc = Document()

    # Set margins to 0.75 inch (standard letter)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    current_table_rows = []

    def clean_text(text):
        # Strip markdown syntax for plain text addition
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        clean = re.sub(r'`(.*?)`', r'\1', clean)
        clean = re.sub(r'\$(.*?)\$', r'\1', clean)
        return clean

    def flush_table():
        nonlocal current_table_rows
        if not current_table_rows:
            return

        table = doc.add_table(rows=len(current_table_rows), cols=len(current_table_rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        for r_idx, row in enumerate(current_table_rows):
            for c_idx, cell_text in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = clean_text(cell_text)
                
                # Header row styling
                if r_idx == 0:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="1A365D"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9)
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9)

        doc.add_paragraph()  # Spacing
        current_table_rows = []

    for line in lines:
        line_str = line.rstrip('\r\n')

        # Code block handling
        if line_str.startswith('```'):
            flush_table()
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Inches(0.2)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(44, 122, 123)
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line_str)
            continue

        # Markdown Table Row Parsing
        if line_str.startswith('|') and '|' in line_str[1:]:
            if '---' in line_str:
                continue
            cols = [c.strip() for c in line_str.split('|')[1:-1]]
            if cols:
                current_table_rows.append(cols)
            continue
        else:
            flush_table()

        if not line_str.strip():
            continue

        # Title
        if line_str.startswith('# '):
            title_text = clean_text(line_str[2:].strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_after = Pt(12)
            continue

        # Section Header (Heading 1 - IEEE Roman Numerals)
        if line_str.startswith('## '):
            h_text = clean_text(line_str[3:].strip())
            p = doc.add_paragraph()
            run = p.add_run(h_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 54, 93)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            continue

        # Subsection Header (Heading 2)
        if line_str.startswith('### '):
            h_text = clean_text(line_str[4:].strip())
            p = doc.add_paragraph()
            run = p.add_run(h_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(43, 108, 176)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            continue

        # Horizontal Rule
        if line_str.strip() == '---':
            continue

        # Bullet List Item
        if line_str.startswith('* ') or line_str.startswith('- '):
            b_text = clean_text(line_str[2:].strip())
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(b_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            continue

        # Normal Paragraph
        p = doc.add_paragraph()
        run = p.add_run(clean_text(line_str))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

    flush_table()
    doc.save(docx_path)
    print(f"Successfully created DOCX paper at: {docx_path}")

if __name__ == '__main__':
    md_path = "/Users/nuwangi/Desktop/research/safe-circle/docs/ICACT2026/SafeCircle_ICACT2026_Paper.md"
    docx_path = "/Users/nuwangi/Desktop/research/safe-circle/docs/ICACT2026/SafeCircle_ICACT2026_Paper.docx"
    create_icact_docx(md_path, docx_path)
